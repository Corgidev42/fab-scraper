import subprocess
import sys

# === Auto-installation des modules requis ===
required_modules = ['selenium', 'beautifulsoup4', 'requests', 'Pillow', 'python-docx']

for module in required_modules:
	try:
		__import__(module if module != 'beautifulsoup4' else 'bs4')
	except ImportError:
		print(f"📦 Module manquant : {module} → installation...")
		subprocess.check_call([sys.executable, "-m", "pip", "install", module])

import os
import time
import requests
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from PIL import Image
from docx import Document

# === Vérification des arguments ===
if len(sys.argv) < 2:
	print("❌ Utilisation : python main.py <URL> [vivid]")
	sys.exit(1)

url = sys.argv[1]

# Option boost couleur
boost_mode = "none"
if len(sys.argv) >= 3:
	mode = sys.argv[2].lower()
	if mode == "vivid":
		boost_mode = "vivid"
		print("🎨 Mode VIVID activé : saturation, contraste et luminosité boostés.")

# === Setup Selenium (Headless propre) ===
delay = 5  # secondes pour attendre le JS
options = Options()
options.add_argument("--headless=new")
options.add_argument("--disable-gpu")
options.add_argument("--window-size=1920,1080")
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")

print("🚀 Lancement de Chrome...")
driver = webdriver.Chrome(options=options)

print(f"🌐 Ouverture de l'URL : {url}")
driver.get(url)
print(f"⏳ Attente de {delay} secondes pour chargement JS...")
time.sleep(delay)

# === Analyse de la page ===
soup = BeautifulSoup(driver.page_source, 'html.parser')
driver.quit()

img_tags = soup.find_all('img', class_='css-1ur2xev')
print(f"📸 {len(img_tags)} image(s) trouvée(s).")

if not img_tags:
	print("❌ Aucune image détectée. Vérifie le chargement ou le sélecteur CSS.")
	sys.exit(1)

# === Dossier basé sur la première carte ===
first_alt = img_tags[0].get('alt', 'deck_sans_nom')
folder_name = first_alt.replace(" ", "_").replace(",", "").replace("'", "").replace("/", "-")
os.makedirs(folder_name, exist_ok=True)
print(f"📂 Dossier de sortie : {folder_name}")

# === Téléchargement des images ===
for img in img_tags:
	img_url = img.get('src')
	alt = img.get('alt', 'image_sans_nom')

	if not img_url:
		continue

	base_filename = alt.replace(" ", "_").replace(",", "").replace("'", "").replace("/", "-")
	extension = img_url.split('.')[-1].split('?')[0]
	filepath = f"{folder_name}/{base_filename}.{extension}"

	counter = 1
	while os.path.exists(filepath):
		filepath = f"{folder_name}/{base_filename}-{counter}.{extension}"
		counter += 1

	try:
		img_data = requests.get(img_url).content
		with open(filepath, 'wb') as f:
			f.write(img_data)
		print(f"✅ Téléchargée : {filepath}")
	except Exception as e:
		print(f"❌ Erreur pour {img_url} : {e}")

# === Traitement impression : CMJN + fond perdu + gabarit fixe ===
def prepare_images_for_print(folder, boost_mode="none"):
	from PIL import ImageEnhance

	dpi = 300
	width_mm, height_mm = 69, 94
	final_width = int(width_mm / 25.4 * dpi)
	final_height = int(height_mm / 25.4 * dpi)
	content_width = int(63 / 25.4 * dpi)
	content_height = int(88 / 25.4 * dpi)
	bleed = int(3 / 25.4 * dpi)

	output_suffix = "_print_ready"
	output_format = "TIFF"

	print(f"\n🖨️  Preparing images for print...")
	print(f"📐 Each image will be resized to: {final_width}x{final_height} px")

	for filename in os.listdir(folder):
		if filename.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
			original_path = os.path.join(folder, filename)
			try:
				img = Image.open(original_path).convert("RGB")

				if boost_mode == "vivid":
					enhancer = ImageEnhance.Color(img)
					img = enhancer.enhance(1.2)
					enhancer = ImageEnhance.Contrast(img)
					img = enhancer.enhance(1.1)
					enhancer = ImageEnhance.Brightness(img)
					img = enhancer.enhance(1.05)

				resized = img.resize((content_width, content_height), Image.LANCZOS)
				final_img = Image.new("RGB", (final_width, final_height), (255, 255, 255))
				final_img.paste(resized, (bleed, bleed))

				cmyk_img = final_img.convert("CMYK")
				output_name = os.path.splitext(filename)[0] + output_suffix + ".tif"
				output_path = os.path.join(folder, output_name)
				cmyk_img.save(output_path, output_format, dpi=(dpi, dpi))
				print(f"🟢 Ready for print: {output_name}")
			except Exception as e:
				print(f"❌ Error processing {filename}: {e}")

prepare_images_for_print(folder_name, boost_mode)

# === Génération PDF final ===
def export_images_to_pdf(folder):
	print(f"\n📄 Création du PDF d'impression...")
	tiff_files = sorted([
		f for f in os.listdir(folder)
		if f.lower().endswith("_print_ready.tif")
	])

	if not tiff_files:
		print("⚠️ Aucune image print-ready trouvée.")
		return

	first_image = Image.open(os.path.join(folder, tiff_files[0]))
	other_images = [Image.open(os.path.join(folder, f)) for f in tiff_files[1:]]

	pdf_path = os.path.join(folder, f"{folder}_print.pdf")
	first_image.save(pdf_path, "PDF", save_all=True, append_images=other_images)
	print(f"✅ PDF final prêt : {pdf_path}")

export_images_to_pdf(folder_name)

# === Génération DOCX liste des cartes ===
def generate_deck_list_doc(folder):
	from docx import Document
	from collections import defaultdict

	card_counts = defaultdict(int)

	for filename in os.listdir(folder):
		if filename.endswith("_print_ready.tif"):
			base = filename.replace("_print_ready.tif", "")
			if "_copy" in base:
				base = base.rsplit("_copy", 1)[0]
			if "-" in base and base.rsplit("-", 1)[-1].isdigit():
				base = base.rsplit("-", 1)[0]
			card_counts[base] += 1

	sorted_cards = sorted(card_counts.items())

	doc = Document()
	doc.add_heading('Deck List', 0)
	for name, qty in sorted_cards:
		doc.add_paragraph(f"{name} x{qty}")

	doc_path = os.path.join(folder, "deck_list.docx")
	doc.save(doc_path)
	print(f"✅ Document Word généré : {doc_path}")

generate_deck_list_doc(folder_name)

# === Demande de suppression des fichiers générés ===
def cleanup_prompt(folder):
	print("Select what you want to delete:")
	print("  t - Print-ready TIFF files")
	print("  i - Original images (webp/jpg/png)")
	print("  p - PDF file")
	print("  d - Word deck list (deck_list.docx)")
	print("  a - All of the above")
	print("  n - Keep everything")

	answer = input("Your choice [t/i/p/d/a/n]: ").strip().lower()

	delete_tiff = 't' in answer or answer == 'a'
	delete_images = 'i' in answer or answer == 'a'
	delete_pdf = 'p' in answer or answer == 'a'
	delete_docx = 'd' in answer or answer == 'a'

	deleted_tif = 0
	deleted_scrap = 0
	deleted_pdf = False
	deleted_doc = False

	if delete_tiff:
		for filename in os.listdir(folder):
			if filename.endswith("_print_ready.tif"):
				try:
					os.remove(os.path.join(folder, filename))
					deleted_tif += 1
				except Exception as e:
					print(f"Error deleting TIFF {filename}: {e}")

	if delete_images:
		for filename in os.listdir(folder):
			if filename.lower().endswith((".webp", ".jpg", ".jpeg", ".png")) and not filename.endswith("_print_ready.tif"):
				try:
					os.remove(os.path.join(folder, filename))
					deleted_scrap += 1
				except Exception as e:
					print(f"Error deleting image {filename}: {e}")

	if delete_pdf:
		pdf_path = os.path.join(folder, f"{folder}_print.pdf")
		if os.path.exists(pdf_path):
			try:
				os.remove(pdf_path)
				deleted_pdf = True
			except Exception as e:
				print(f"Error deleting PDF: {e}")

	if delete_docx:
		doc_path = os.path.join(folder, "deck_list.docx")
		if os.path.exists(doc_path):
			try:
				os.remove(doc_path)
				deleted_doc = True
			except Exception as e:
				print(f"Error deleting Word file: {e}")

	print()
	if answer == 'n':
		print("No files deleted.")
	else:
		if delete_tiff:
			print(f"Deleted {deleted_tif} TIFF file(s).")
		if delete_images:
			print(f"Deleted {deleted_scrap} original image(s).")
		if delete_pdf and deleted_pdf:
			print("Deleted PDF file.")
		if delete_docx and deleted_doc:
			print("Deleted Word file.")

# Appel final
cleanup_prompt(folder_name)
